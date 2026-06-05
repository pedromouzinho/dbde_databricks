# =============================================================================
# export_engine.py — Export CSV/XLSX/PDF/SVG/HTML v7.0
# =============================================================================
# Gera ficheiros de export a partir de dados de tool results.
# CSV/XLSX: stdlib + openpyxl. PDF: fpdf2. SVG: geração manual.
# =============================================================================

import csv
import io
import json
import logging
import html as html_lib
import os
import re
import zipfile
from datetime import datetime, timezone
from typing import Dict, List, Any, Optional
from urllib.parse import urlparse

from config_databricks import EXPORT_BRAND_COLOR, EXPORT_BRAND_NAME, EXPORT_AGENT_NAME, APP_VERSION

logger = logging.getLogger(__name__)

# =============================================================================
# DATA EXTRACTION
# =============================================================================

def extract_table_data(tool_result: dict) -> tuple[List[str], List[List[str]]]:
    """Extrai headers e rows de um tool result."""
    items = tool_result.get("items", tool_result.get("analysis_data", []))
    if not items:
        # KPI groups?
        groups = tool_result.get("groups", [])
        if groups:
            return ["Valor", "Contagem"], [[g["value"], str(g["count"])] for g in groups]
        # Timeline?
        timeline = tool_result.get("timeline", [])
        if timeline:
            return ["Mês", "Contagem"], [[t[0], str(t[1])] for t in timeline]
        return [], []
    
    # Determinar headers a partir das keys do primeiro item
    sample = items[0]
    # Ordem preferencial
    preferred = ["id", "type", "title", "state", "area", "assigned_to", "created_by", "created_date", "url"]
    headers = [k for k in preferred if k in sample]
    headers.extend(k for k in sample.keys() if k not in headers and k != "score")
    
    rows = []
    for item in items:
        rows.append([str(item.get(h, "")) for h in headers])
    
    return headers, rows


def _clean_header(h: str) -> str:
    """Limpa header para display."""
    return h.replace("_", " ").title()


def _safe_sheet_title(title: str) -> str:
    """Sanitize Excel worksheet title (max 31 chars, no []:*?/\\)."""
    if not title:
        return "Export"
    forbidden = set('[]:*?/\\')
    safe = ''.join('_' if ch in forbidden else ch for ch in str(title))
    safe = safe.strip().strip("'")
    if not safe:
        safe = "Export"
    return safe[:31]


def _latin1_safe(text: str, max_len: int = 0) -> str:
    """Sanitize text for fpdf2 core fonts (Latin-1 only)."""
    if not text:
        return ""
    import unicodedata
    nfkd = unicodedata.normalize("NFKD", text)
    cleaned = "".join(ch for ch in nfkd if not unicodedata.combining(ch))
    cleaned = (
        cleaned.replace("—", "-")
        .replace("–", "-")
        .replace("“", '"')
        .replace("”", '"')
        .replace("’", "'")
    )
    safe = cleaned.encode("latin-1", errors="replace").decode("latin-1")
    if max_len > 0:
        safe = safe[:max_len]
    return safe


def _safe_http_url(value: str) -> str:
    candidate = str(value or "").strip()
    if not candidate:
        return ""
    try:
        parsed = urlparse(candidate)
    except Exception:
        return ""
    if parsed.scheme in ("http", "https") and parsed.netloc:
        return candidate
    return ""


def _hex_to_rgb(hex_color: str, fallback: tuple[int, int, int]) -> tuple[int, int, int]:
    raw = str(hex_color or "").strip().lstrip("#")
    if len(raw) != 6:
        return fallback
    try:
        return (int(raw[0:2], 16), int(raw[2:4], 16), int(raw[4:6], 16))
    except Exception:
        return fallback


def _is_numeric_like(value: str) -> bool:
    txt = str(value or "").strip()
    if not txt:
        return False
    return bool(re.fullmatch(r"[+-]?\d+(?:[.,]\d+)?", txt))


def _configure_pdf_font(pdf) -> str:
    """Configura Montserrat quando existir no runtime; fallback para Helvetica."""
    base_dir = os.path.dirname(os.path.abspath(__file__))
    font_dir = os.path.join(base_dir, "assets", "fonts")
    regular = os.path.join(font_dir, "Montserrat-Regular.ttf")
    bold = os.path.join(font_dir, "Montserrat-Bold.ttf")
    italic = os.path.join(font_dir, "Montserrat-Italic.ttf")
    if os.path.exists(regular) and os.path.exists(bold) and os.path.exists(italic):
        try:
            pdf.add_font("Montserrat", "", regular)
            pdf.add_font("Montserrat", "B", bold)
            pdf.add_font("Montserrat", "I", italic)
            return "Montserrat"
        except Exception as e:
            logging.warning("[ExportEngine] Montserrat unavailable, fallback Helvetica: %s", e)
    return "Helvetica"


def _pdf_column_widths(headers: List[str], page_width: float) -> List[float]:
    if not headers:
        return []
    weights = {
        "id": 1.0,
        "type": 1.2,
        "state": 1.1,
        "parent_id": 1.1,
        "title": 4.8,
        "area": 2.6,
        "assigned_to": 1.9,
        "created_by": 1.9,
        "created_date": 1.6,
        "url": 3.0,
    }
    min_w = 14.0
    raw = [float(weights.get(h, 1.6)) for h in headers]
    total = sum(raw) if raw else 1.0
    widths = [max(min_w, page_width * (w / total)) for w in raw]
    overflow = sum(widths) - page_width
    if overflow > 0:
        candidates = sorted(range(len(widths)), key=lambda idx: widths[idx], reverse=True)
        for idx in candidates:
            if overflow <= 0:
                break
            shrink = max(0.0, widths[idx] - min_w)
            if shrink <= 0:
                continue
            delta = min(shrink, overflow)
            widths[idx] -= delta
            overflow -= delta
    if sum(widths) > page_width and headers:
        eq = max(10.0, page_width / len(headers))
        widths = [eq] * len(headers)
    return widths


def _pdf_wrap_lines(pdf, text: str, cell_width: float, max_lines: int = 4) -> List[str]:
    raw = _latin1_safe(str(text or "")).strip()
    if not raw:
        return [""]

    max_text_w = max(2.0, cell_width - 1.8)
    lines: List[str] = []

    for segment in raw.replace("\r", "\n").split("\n"):
        words = segment.split()
        if not words:
            continue
        current = words[0]
        if pdf.get_string_width(current) > max_text_w:
            # Hard-break de tokens longos (URL, IDs) para manter largura estável.
            chunk = ""
            for ch in current:
                cand = chunk + ch
                if pdf.get_string_width(cand) <= max_text_w:
                    chunk = cand
                else:
                    if chunk:
                        lines.append(chunk)
                    chunk = ch
            current = chunk or current
        for word in words[1:]:
            if pdf.get_string_width(word) > max_text_w:
                lines.append(current)
                chunk = ""
                for ch in word:
                    cand = chunk + ch
                    if pdf.get_string_width(cand) <= max_text_w:
                        chunk = cand
                    else:
                        if chunk:
                            lines.append(chunk)
                        chunk = ch
                current = chunk or word
                continue
            candidate = f"{current} {word}".strip()
            if pdf.get_string_width(candidate) <= max_text_w:
                current = candidate
            else:
                lines.append(current)
                current = word
        lines.append(current)

    if len(lines) > max_lines:
        lines = lines[:max_lines]
        if lines:
            tail = lines[-1]
            if len(tail) > 2:
                lines[-1] = tail[:-1] + "..."
            else:
                lines[-1] = tail + "..."
    return [ln for ln in lines if ln] or [""]


# =============================================================================
# CSV EXPORT
# =============================================================================

def to_csv(tool_result: dict, filename: str = "export.csv") -> io.BytesIO:
    """Gera CSV (UTF-8 BOM para compatibilidade Excel)."""
    headers, rows = extract_table_data(tool_result)
    
    buf = io.BytesIO()
    buf.write(b'\xef\xbb\xbf')  # UTF-8 BOM
    
    wrapper = io.TextIOWrapper(buf, encoding='utf-8', newline='')
    writer = csv.writer(wrapper)
    writer.writerow([_clean_header(h) for h in headers])
    writer.writerows(rows)
    wrapper.flush()
    wrapper.detach()
    
    buf.seek(0)
    return buf


# =============================================================================
# XLSX EXPORT (openpyxl)
# =============================================================================

def to_xlsx(tool_result: dict, title: str = "Export", filename: str = "export.xlsx") -> io.BytesIO:
    """Gera XLSX formatado com branding Millennium."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        # Fallback to CSV if openpyxl not available
        return to_csv(tool_result, filename.replace(".xlsx", ".csv"))
    
    headers, rows = extract_table_data(tool_result)
    if not headers:
        return to_csv(tool_result)
    
    wb = Workbook()
    ws = wb.active
    ws.title = _safe_sheet_title(title)
    
    # Branding colors
    brand_hex = str(EXPORT_BRAND_COLOR or "#DE3163").strip().lstrip("#").upper()
    if len(brand_hex) != 6:
        brand_hex = "DE3163"
    brand_fill = PatternFill(start_color=brand_hex, end_color=brand_hex, fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    data_font = Font(size=10)
    zebra_fill = PatternFill(start_color="F5F5F5", end_color="F5F5F5", fill_type="solid")
    border = Border(
        left=Side(style='thin', color='DDDDDD'),
        right=Side(style='thin', color='DDDDDD'),
        top=Side(style='thin', color='DDDDDD'),
        bottom=Side(style='thin', color='DDDDDD'),
    )
    
    # Title row
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
    title_cell = ws.cell(row=1, column=1, value=f"{EXPORT_AGENT_NAME} — {title}")
    title_cell.font = Font(bold=True, size=14, color=brand_hex)
    
    # Subtitle
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(headers))
    # Local time intentional for user-facing display.
    sub = ws.cell(row=2, column=1, value=f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')} | Total: {tool_result.get('total_count', len(rows))} registos")
    sub.font = Font(size=9, color="666666", italic=True)
    
    # Headers (row 4)
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=4, column=col, value=_clean_header(h))
        cell.font = header_font
        cell.fill = brand_fill
        cell.alignment = Alignment(horizontal='center')
        cell.border = border
    
    # Data rows
    for r_idx, row in enumerate(rows, 5):
        for c_idx, val in enumerate(row, 1):
            cell = ws.cell(row=r_idx, column=c_idx, value=val)
            cell.font = data_font
            cell.border = border
            if (r_idx - 5) % 2 == 1:
                cell.fill = zebra_fill
    
    # Auto-width
    for col in range(1, len(headers) + 1):
        max_len = len(_clean_header(headers[col-1]))
        for row in rows[:50]:  # Sample for performance
            if col-1 < len(row):
                max_len = max(max_len, len(str(row[col-1])))
        ws.column_dimensions[get_column_letter(col)].width = min(max_len + 4, 50)
    
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def _fallback_docx_from_lines(lines: List[str]) -> io.BytesIO:
    """Gera DOCX mínimo sem dependências externas (fallback)."""
    def _xml_escape(value: str) -> str:
        return html_lib.escape(str(value or ""), quote=False)

    body_xml = "".join(
        f"<w:p><w:r><w:t>{_xml_escape(line)}</w:t></w:r></w:p>"
        for line in lines
    )
    if not body_xml:
        body_xml = "<w:p><w:r><w:t>Sem dados para exportar.</w:t></w:r></w:p>"

    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" '
        'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
        'xmlns:o="urn:schemas-microsoft-com:office:office" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
        'xmlns:v="urn:schemas-microsoft-com:vml" '
        'xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" '
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:w10="urn:schemas-microsoft-com:office:word" '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" '
        'xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" '
        'xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk" '
        'xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml" '
        'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" '
        'mc:Ignorable="w14 wp14">'
        f"<w:body>{body_xml}<w:sectPr/></w:body></w:document>"
    )

    content_types = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        '</Types>'
    )
    rels = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
        'Target="word/document.xml"/>'
        '</Relationships>'
    )
    doc_rels = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"></Relationships>'
    )

    out = io.BytesIO()
    with zipfile.ZipFile(out, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", rels)
        zf.writestr("word/document.xml", document_xml)
        zf.writestr("word/_rels/document.xml.rels", doc_rels)
    out.seek(0)
    return out


def to_docx(tool_result: dict, title: str = "Export", filename: str = "export.docx") -> io.BytesIO:
    """Gera DOCX com branding e tabela simples."""
    headers, rows = extract_table_data(tool_result)
    generated_at = datetime.now().strftime("%d/%m/%Y %H:%M")
    total = tool_result.get("total_count", len(rows))

    try:
        from docx import Document

        doc = Document()
        doc.add_heading(f"{EXPORT_AGENT_NAME} — {title}", level=1)
        doc.add_paragraph(f"Gerado em {generated_at} | Total: {total} registos")

        if headers:
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            for idx, header in enumerate(headers):
                run = hdr_cells[idx].paragraphs[0].add_run(_clean_header(header))
                run.bold = True
            for row in rows:
                row_cells = table.add_row().cells
                for idx, val in enumerate(row):
                    row_cells[idx].text = str(val)
        else:
            doc.add_paragraph("Sem dados para exportar.")

        doc.add_paragraph(f"{EXPORT_BRAND_NAME} | {EXPORT_AGENT_NAME} v{APP_VERSION}")
        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        return out
    except Exception as e:
        logging.warning("[ExportEngine] to_docx fallback ativo: %s", e)
        lines = [f"{EXPORT_AGENT_NAME} — {title}", f"Gerado em {generated_at} | Total: {total} registos"]
        if headers and rows:
            lines.append(" | ".join(_clean_header(h) for h in headers))
            lines.extend(" | ".join(str(v) for v in row) for row in rows)
        else:
            lines.append("Sem dados para exportar.")
        lines.append(f"{EXPORT_BRAND_NAME} | {EXPORT_AGENT_NAME} v{APP_VERSION}")
        return _fallback_docx_from_lines(lines)


# =============================================================================
# PDF EXPORT (fpdf2)
# =============================================================================

def to_pdf(tool_result: dict, title: str = "Export", summary: str = "") -> io.BytesIO:
    """Gera PDF com tabela de dados."""
    try:
        from fpdf import FPDF
        from fpdf.enums import XPos, YPos
    except ImportError:
        # Fallback
        buf = io.BytesIO()
        buf.write(b"PDF generation requires fpdf2. Install: pip install fpdf2")
        buf.seek(0)
        return buf

    try:
        headers, rows = extract_table_data(tool_result)

        pdf = FPDF()
        pdf.add_page('L')  # Landscape for tables
        pdf.set_auto_page_break(auto=True, margin=15)
        font_family = _configure_pdf_font(pdf)
        cerise_rgb = _hex_to_rgb(str(EXPORT_BRAND_COLOR or "#DE3163"), (222, 49, 99))

        # Title
        pdf.set_font(font_family, 'B', 16)
        pdf.set_text_color(*cerise_rgb)
        pdf.cell(0, 10, _latin1_safe(title, 160), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        # Subtitle
        pdf.set_font(font_family, '', 9)
        pdf.set_text_color(100, 100, 100)
        # Local time intentional for user-facing display.
        pdf.multi_cell(0, 5, _latin1_safe(f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')} | {EXPORT_AGENT_NAME}"))
        if summary:
            pdf.set_font(font_family, 'I', 8.5)
            for line in str(summary).splitlines():
                if line.strip():
                    try:
                        pdf.multi_cell(0, 4.5, _latin1_safe(line))
                    except Exception:
                        pdf.cell(0, 4.5, _latin1_safe(line, 250), ln=True)
        pdf.ln(5)

        if not headers:
            pdf.set_font(font_family, '', 10)
            pdf.cell(0, 10, "Sem dados para exportar.", ln=True)
        else:
            page_w = pdf.w - pdf.l_margin - pdf.r_margin
            col_widths = _pdf_column_widths(headers, page_w)

            def _draw_header() -> None:
                pdf.set_font(font_family, 'B', 8)
                pdf.set_fill_color(*cerise_rgb)
                pdf.set_text_color(255, 255, 255)
                pdf.set_x(pdf.l_margin)
                for i, h in enumerate(headers):
                    pdf.cell(col_widths[i], 7, _latin1_safe(_clean_header(h), 28), border=1, fill=True)
                pdf.ln()

            _draw_header()

            pdf.set_font(font_family, '', 7)
            pdf.set_text_color(0, 0, 0)
            line_h = 3.8

            prepared_rows: List[tuple[List[str], List[List[str]]]] = []
            global_max_lines = 1
            for row in rows[:500]:
                cell_lines: List[List[str]] = []
                for i, val in enumerate(row):
                    key = headers[i] if i < len(headers) else ""
                    max_lines = 6 if key in ("title", "url", "area") else 4
                    wrapped = _pdf_wrap_lines(pdf, str(val), col_widths[i], max_lines=max_lines)
                    cell_lines.append(wrapped)
                    global_max_lines = max(global_max_lines, len(wrapped))
                prepared_rows.append((row, cell_lines))

            uniform_row_lines = max(2, min(6, global_max_lines))

            for r_idx, (row_values, cell_lines) in enumerate(prepared_rows):
                row_h = line_h * uniform_row_lines
                if pdf.get_y() + row_h > (pdf.h - pdf.b_margin - 10):
                    pdf.add_page('L')
                    _draw_header()

                fill = (r_idx % 2 == 1)
                if fill:
                    pdf.set_fill_color(245, 245, 245)

                y0 = pdf.get_y()
                x = pdf.l_margin
                for i, lines in enumerate(cell_lines):
                    padded_lines = (lines + ([""] * max(0, uniform_row_lines - len(lines))))[:uniform_row_lines]
                    cell_style = "I" if _is_numeric_like(row_values[i] if i < len(row_values) else "") else ""
                    pdf.set_font(font_family, cell_style, 7)
                    pdf.set_xy(x, y0)
                    cell_text = "\n".join(padded_lines)
                    try:
                        pdf.multi_cell(
                            col_widths[i],
                            line_h,
                            cell_text,
                            border=1,
                            fill=fill,
                        )
                    except Exception:
                        pdf.set_xy(x, y0)
                        pdf.cell(
                            col_widths[i],
                            row_h,
                            _latin1_safe(cell_text.replace("\n", " | "), 220),
                            border=1,
                            fill=fill,
                        )
                    x += col_widths[i]
                pdf.set_xy(pdf.l_margin, y0 + row_h)

        # Footer
        pdf.ln(10)
        pdf.set_font(font_family, 'I', 8)
        pdf.set_text_color(150, 150, 150)
        pdf.cell(
            0,
            5,
            _latin1_safe(f"{EXPORT_BRAND_NAME} | {EXPORT_AGENT_NAME} v{APP_VERSION}"),
            new_x=XPos.LMARGIN,
            new_y=YPos.NEXT,
        )

        buf = io.BytesIO()
        pdf.output(buf)
        buf.seek(0)
        return buf
    except Exception as e:
        logging.error("[ExportEngine] to_pdf failed: %s", e)
        pdf_fallback = FPDF()
        pdf_fallback.add_page()
        pdf_fallback.set_font('Helvetica', '', 12)
        pdf_fallback.cell(
            0,
            10,
            "Erro ao gerar PDF. Tenta CSV ou XLSX.",
            new_x=XPos.LMARGIN,
            new_y=YPos.NEXT,
        )
        buf = io.BytesIO()
        pdf_fallback.output(buf)
        buf.seek(0)
        return buf


# =============================================================================
# SVG CHART EXPORT
# =============================================================================

def to_svg_bar_chart(tool_result: dict, title: str = "Chart") -> str:
    """Gera SVG bar chart simples a partir de groups ou distribution."""
    groups = tool_result.get("groups", [])
    if not groups:
        # Try state_distribution
        dist = tool_result.get("state_distribution", tool_result.get("type_distribution", {}))
        if dist:
            groups = [{"value": k, "count": v} for k, v in sorted(dist.items(), key=lambda x: x[1], reverse=True)]
    
    if not groups:
        return '<svg xmlns="http://www.w3.org/2000/svg" width="400" height="100"><text x="20" y="50">Sem dados</text></svg>'
    
    groups = groups[:20]  # Max 20 bars
    max_val = max(g["count"] for g in groups) or 1
    
    bar_h = 28
    label_w = 180
    chart_w = 600
    bar_w = chart_w - label_w - 60
    total_h = len(groups) * bar_h + 80
    
    svg = [f'<svg xmlns="http://www.w3.org/2000/svg" width="{chart_w}" height="{total_h}" style="font-family:Arial,sans-serif">']
    
    # Title
    safe_title = html_lib.escape(str(title or "Chart"))
    svg.append(f'<text x="{chart_w//2}" y="25" text-anchor="middle" font-size="16" font-weight="bold" fill="{html_lib.escape(EXPORT_BRAND_COLOR)}">{safe_title}</text>')
    
    y = 50
    for g in groups:
        w = int((g["count"] / max_val) * bar_w)
        label = html_lib.escape(str(g["value"])[:25])
        
        svg.append(f'<text x="{label_w - 10}" y="{y + 16}" text-anchor="end" font-size="11" fill="#333">{label}</text>')
        svg.append(f'<rect x="{label_w}" y="{y + 2}" width="{max(w, 2)}" height="{bar_h - 6}" fill="{html_lib.escape(EXPORT_BRAND_COLOR)}" rx="3"/>')
        svg.append(f'<text x="{label_w + w + 8}" y="{y + 16}" font-size="11" fill="#666">{g["count"]}</text>')
        
        y += bar_h
    
    # Footer
    # Local time intentional for user-facing display.
    svg.append(f'<text x="{chart_w//2}" y="{total_h - 10}" text-anchor="middle" font-size="9" fill="#999">{EXPORT_AGENT_NAME} | {datetime.now().strftime("%d/%m/%Y")}</text>')
    svg.append('</svg>')
    
    return '\n'.join(svg)


# =============================================================================
# HTML REPORT EXPORT
# =============================================================================

def to_html_report(tool_result: dict, title: str = "Relatório", summary: str = "") -> str:
    """Gera relatório HTML completo com tabela e estilos."""
    headers, rows = extract_table_data(tool_result)
    safe_title = html_lib.escape(str(title or "Relatório"))
    safe_summary = html_lib.escape(str(summary or ""))
    total_count = html_lib.escape(str(tool_result.get('total_count', len(rows))))
    # Local time intentional for user-facing display.
    generated_at = html_lib.escape(datetime.now().strftime('%d/%m/%Y %H:%M'))
    
    html = f"""<!DOCTYPE html>
<html lang="pt">
<head><meta charset="UTF-8"><title>{safe_title}</title>
<style>
body {{ font-family: 'Segoe UI', Arial, sans-serif; margin: 40px; color: #333; }}
h1 {{ color: {html_lib.escape(EXPORT_BRAND_COLOR)}; border-bottom: 3px solid {html_lib.escape(EXPORT_BRAND_COLOR)}; padding-bottom: 10px; }}
.meta {{ color: #666; font-size: 0.9em; margin-bottom: 20px; }}
table {{ border-collapse: collapse; width: 100%; margin-top: 20px; }}
th {{ background: {html_lib.escape(EXPORT_BRAND_COLOR)}; color: white; padding: 10px 12px; text-align: left; font-size: 0.9em; }}
td {{ padding: 8px 12px; border-bottom: 1px solid #eee; font-size: 0.85em; }}
tr:nth-child(even) {{ background: #f9f9f9; }}
tr:hover {{ background: #fff0f0; }}
.footer {{ margin-top: 30px; font-size: 0.8em; color: #999; border-top: 1px solid #eee; padding-top: 10px; }}
a {{ color: {html_lib.escape(EXPORT_BRAND_COLOR)}; text-decoration: none; }}
</style></head>
<body>
<h1>{safe_title}</h1>
<div class="meta">Gerado em {generated_at} | Total: {total_count} registos{f' | {safe_summary}' if safe_summary else ''}</div>
"""
    if headers:
        html += '<table><thead><tr>'
        for h in headers:
            html += f'<th>{html_lib.escape(_clean_header(h))}</th>'
        html += '</tr></thead><tbody>'
        for row in rows:
            html += '<tr>'
            for i, val in enumerate(row):
                text_val = html_lib.escape(str(val))
                if headers[i] == 'url':
                    safe_url = _safe_http_url(str(val))
                    if safe_url:
                        escaped_url = html_lib.escape(safe_url, quote=True)
                        html += f'<td><a href="{escaped_url}" target="_blank" rel="noopener noreferrer">🔗 Link</a></td>'
                    else:
                        html += f'<td>{text_val}</td>'
                elif headers[i] == 'id' and len(row) > headers.index('url') if 'url' in headers else False:
                    url_val = row[headers.index('url')] if 'url' in headers else ''
                    safe_url = _safe_http_url(str(url_val))
                    if safe_url:
                        escaped_url = html_lib.escape(safe_url, quote=True)
                        html += f'<td><a href="{escaped_url}" target="_blank" rel="noopener noreferrer">{text_val}</a></td>'
                    else:
                        html += f'<td>{text_val}</td>'
                else:
                    html += f'<td>{text_val}</td>'
            html += '</tr>'
        html += '</tbody></table>'
    else:
        html += '<p>Sem dados tabulares para apresentar.</p>'
        # Show raw data
        raw = json.dumps(tool_result, indent=2, ensure_ascii=False)[:5000]
        html += f'<pre>{html_lib.escape(raw)}</pre>'
    
    html += f'\n<div class="footer">{EXPORT_BRAND_NAME} | {EXPORT_AGENT_NAME} v{APP_VERSION}</div>\n</body></html>'
    return html


def to_html(tool_result: dict, title: str = "Export") -> io.BytesIO:
    """Gera HTML em BytesIO para paridade com to_csv/to_xlsx/to_pdf/to_docx."""
    html_str = to_html_report(tool_result, title=title)
    buf = io.BytesIO()
    buf.write(html_str.encode("utf-8"))
    buf.seek(0)
    return buf


# =============================================================================
# CHAT PDF EXPORT (fpdf2) — substitui weasyprint
# =============================================================================

def to_chat_pdf(messages: list, title: str = "Chat Export") -> bytes:
    """Gera PDF de conversa de chat com fpdf2. Substitui weasyprint.

    Args:
        messages: lista de dicts com 'role' e 'content'.
        title: titulo do export.
    Returns:
        bytes do PDF gerado.
    """
    try:
        from fpdf import FPDF
    except ImportError:
        raise RuntimeError("fpdf2 required for chat PDF export")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    font_family = _configure_pdf_font(pdf)
    cerise_rgb = _hex_to_rgb(str(EXPORT_BRAND_COLOR or "#DE3163"), (222, 49, 99))

    pdf.add_page()

    # --- Header ---
    pdf.set_font(font_family, "B", 16)
    pdf.set_text_color(*cerise_rgb)
    pdf.cell(0, 10, _latin1_safe(title, 120), ln=True)

    pdf.set_font(font_family, "", 9)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(
        0, 5,
        _latin1_safe(f"Exportado em {datetime.now().strftime('%d/%m/%Y %H:%M')} | {EXPORT_AGENT_NAME}"),
        ln=True,
    )
    pdf.ln(6)

    # --- Messages ---
    page_w = pdf.w - pdf.l_margin - pdf.r_margin

    for msg in messages:
        role = str(msg.get("role", "user")).strip().lower()
        content = str(msg.get("content", "")).strip()
        if not content:
            continue

        # Role label
        if role == "assistant":
            label = "Assistente"
            label_rgb = cerise_rgb
        elif role == "system":
            label = "Sistema"
            label_rgb = (100, 100, 100)
        else:
            label = "Utilizador"
            label_rgb = (50, 50, 50)

        # Check page space — add page if < 30mm left
        if pdf.get_y() > pdf.h - 30:
            pdf.add_page()

        # Role badge
        pdf.set_font(font_family, "B", 9)
        pdf.set_text_color(*label_rgb)
        pdf.cell(0, 5, _latin1_safe(label), ln=True)

        # Content
        pdf.set_font(font_family, "", 9)
        pdf.set_text_color(30, 30, 30)

        safe_content = _latin1_safe(content, 20000)
        for paragraph in safe_content.split("\n"):
            paragraph = paragraph.strip()
            if not paragraph:
                pdf.ln(2)
                continue
            try:
                pdf.multi_cell(page_w, 4.5, paragraph)
            except Exception:
                # Fallback for problematic text
                pdf.cell(0, 4.5, paragraph[:500], ln=True)

        # Separator line
        pdf.ln(3)
        y = pdf.get_y()
        pdf.set_draw_color(220, 220, 220)
        pdf.line(pdf.l_margin, y, pdf.l_margin + page_w, y)
        pdf.ln(4)

    # --- Footer ---
    pdf.ln(5)
    pdf.set_font(font_family, "I", 8)
    pdf.set_text_color(150, 150, 150)
    pdf.cell(
        0, 5,
        _latin1_safe(f"{EXPORT_BRAND_NAME} | {EXPORT_AGENT_NAME} v{APP_VERSION}"),
        ln=True,
    )

    buf = io.BytesIO()
    pdf.output(buf)
    return buf.getvalue()
